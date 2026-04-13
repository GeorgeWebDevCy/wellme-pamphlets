#!/usr/bin/env python
"""
Build a WELLME module import package from the local project source docs.

The resulting ZIP is designed for the plugin's admin-side importer and contains:
  - manifest.json
  - module media assets referenced by the manifest
"""

from __future__ import annotations

import argparse
import html
import json
import os
import re
import shutil
import tempfile
import zipfile
from collections import OrderedDict
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable, Iterator

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from PIL import Image


DEFAULT_SOURCE_ROOT = Path(
    r"C:\Users\georg\Downloads\WellMe-20260413T092223Z-3-001\WellMe"
)
WP2_RELATIVE_DIR = Path("WP2.Measuring Youth Wellbeing-Learning Modules") / "WP2.4_Design_Learning_Modules"
DEFAULT_OUTPUT = Path("dist") / "wellme-import-package.zip"

SECTION_MARKERS = OrderedDict(
    [
        ("INTRODUCTION TO MODULE", "introduction"),
        ("INTRODUCTION", "introduction"),
        ("OVERVIEW", "overview"),
        ("GOALS & OBJECTIVES", "goals"),
        ("GOALS AND OBJECTIVES", "goals"),
        ("THEORETICAL BACKGROUND", "theory"),
        ("FROM THEORY TO PRACTICE", "practice"),
        ("ASSESSMENT", "assessment"),
        ("CORRECT ANSWERS", "answers"),
        ("ANNEX 1", "annex"),
        ("ANNEX", "annex"),
        ("REFERENCES", "references"),
    ]
)

SECTION_LABELS = {
    "introduction": "Introduction",
    "overview": "Overview",
    "goals": "Goals & Objectives",
    "theory": "Theoretical Background",
    "practice": "From Theory to Practice",
    "assessment": "Assessment",
    "answers": "Correct Answers",
    "annex": "Annex",
    "references": "References",
}

MODULE_COLORS = {
    1: "#1D4ED8",
    2: "#0F766E",
    3: "#16A34A",
    4: "#F97316",
    5: "#EAB308",
    6: "#7C3AED",
}

MODULE_SUBTITLES = {
    1: "Strength-based practice",
    2: "Resilience and adaptability",
    3: "Healthy routines",
    4: "Move for wellbeing",
    5: "Spaces for belonging",
    6: "Community-powered transitions",
}

MODULE_MOTTOS = {
    1: "Strengths grow where young people are seen, trusted, and supported.",
    2: "Resilience grows through reflection, adaptation, and practice.",
    3: "Healthy nutrition fuels wellbeing, energy, and participation.",
    4: "Movement turns wellbeing into everyday action.",
    5: "Belonging begins in spaces designed with care and connection.",
    6: "Strong transitions are built through community, guidance, and agency.",
}


@dataclass
class SectionData:
    key: str
    blocks: list[str] = field(default_factory=list)
    texts: list[str] = field(default_factory=list)
    exercises: list[dict] = field(default_factory=list)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build a WELLME import package ZIP.")
    parser.add_argument(
        "--source-root",
        type=Path,
        default=DEFAULT_SOURCE_ROOT,
        help="Root WellMe directory that contains the work-package folders.",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=DEFAULT_OUTPUT,
        help="Destination ZIP path for the generated import package.",
    )
    return parser.parse_args()


def normalize_text(text: str) -> str:
    return " ".join(text.replace("\xa0", " ").split())


def slugify(value: str) -> str:
    value = normalize_text(value).lower()
    value = re.sub(r"[^a-z0-9]+", "-", value)
    return value.strip("-")


def escape_paragraph(text: str) -> str:
    return html.escape(normalize_text(text), quote=False)


def paragraph_is_list(paragraph: Paragraph) -> bool:
    if paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None:
        return True
    stripped = paragraph.text.strip()
    return stripped.startswith(("- ", "• ", "* "))


def paragraph_as_heading(paragraph: Paragraph, text: str) -> bool:
    style_name = (paragraph.style.name or "").lower()
    if not text:
        return False
    if style_name.startswith("heading"):
        return len(text) <= 140 and not text.endswith(".")
    if text.isupper() and len(text) <= 100:
        return True
    return False


def paragraph_html(paragraph: Paragraph) -> str:
    text = normalize_text(paragraph.text)
    if not text:
        return ""
    if paragraph_as_heading(paragraph, text):
        return f"<h3>{html.escape(text, quote=False)}</h3>"
    return f"<p>{html.escape(text, quote=False)}</p>"


def table_html(table: Table) -> str:
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_text = "<br>".join(
                html.escape(normalize_text(p.text), quote=False)
                for p in cell.paragraphs
                if normalize_text(p.text)
            )
            cells.append(f"<td>{cell_text}</td>")
        if cells:
            rows.append("<tr>" + "".join(cells) + "</tr>")
    if not rows:
        return ""
    return '<table class="wellme-import-table">' + "".join(rows) + "</table>"


def iter_block_items(parent: DocxDocument | _Cell) -> Iterator[Paragraph | Table]:
    if isinstance(parent, DocxDocument):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise TypeError(f"Unsupported parent type: {type(parent)!r}")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def detect_section_marker(text: str) -> tuple[str | None, str]:
    upper_text = normalize_text(text).upper()
    for marker, section_key in SECTION_MARKERS.items():
        if upper_text == marker:
            return section_key, ""
        if upper_text.startswith(marker + " "):
            remainder = text[len(marker) :].strip(" :-")
            return section_key, remainder
    return None, ""


def extract_module_number(path: Path) -> int:
    match = re.search(r"module[_\s.-]*no[._\s-]*(\d)", path.name, flags=re.IGNORECASE)
    if match:
        return int(match.group(1))
    match = re.match(r"(\d)\.", path.name)
    if match:
        return int(match.group(1))
    raise ValueError(f"Could not determine module number from {path.name}")


def find_title(doc: Document) -> str:
    for paragraph in doc.paragraphs:
        text = normalize_text(paragraph.text)
        if not text:
            continue
        style_name = (paragraph.style.name or "").lower()
        if style_name.startswith("heading 1"):
            return text

    for paragraph in doc.paragraphs:
        text = normalize_text(paragraph.text)
        if not text:
            continue
        if text.upper().startswith("MODULE "):
            continue
        if text.lower().startswith("introduction to module"):
            continue
        return text

    raise ValueError("Could not locate a usable module title")


def finalize_exercise(current_exercise: dict | None, section: SectionData | None) -> None:
    if not current_exercise or not section:
        return
    content = "".join(current_exercise["blocks"]).strip()
    if content:
        section.exercises.append(
            {
                "title": current_exercise["title"],
                "content": content,
                "texts": current_exercise["texts"],
            }
        )


def parse_wp2_doc(path: Path) -> dict:
    doc = Document(path)
    title = find_title(doc)
    module_number = extract_module_number(path)
    sections: OrderedDict[str, SectionData] = OrderedDict()
    current_section: SectionData | None = None
    current_exercise: dict | None = None

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = normalize_text(block.text)
            if not text:
                continue
            if text.upper().startswith(f"MODULE {module_number}"):
                continue
            if text == title:
                continue

            section_key, remainder = detect_section_marker(text)
            if section_key:
                finalize_exercise(current_exercise, current_section)
                current_exercise = None
                current_section = sections.setdefault(section_key, SectionData(key=section_key))
                if remainder:
                    current_section.blocks.append(f"<p>{html.escape(remainder, quote=False)}</p>")
                    current_section.texts.append(remainder)
                continue

            if current_section is None:
                continue

            if current_section.key == "practice":
                exercise_match = re.match(r"Exercise\s+\d+[:\-]?\s*(.+)", text, flags=re.IGNORECASE)
                if exercise_match:
                    finalize_exercise(current_exercise, current_section)
                    current_exercise = {
                        "title": normalize_text(exercise_match.group(1) or text),
                        "blocks": [],
                        "texts": [],
                    }
                    continue

            if paragraph_is_list(block):
                current_html = f"<li>{escape_paragraph(text.lstrip('-*• '))}</li>"
                if current_exercise is not None:
                    current_exercise["blocks"].append(current_html)
                    current_exercise["texts"].append(text)
                else:
                    current_section.blocks.append(current_html)
                    current_section.texts.append(text)
                continue

            rendered = paragraph_html(block)
            if not rendered:
                continue

            if current_exercise is not None:
                current_exercise["blocks"].append(rendered)
                current_exercise["texts"].append(text)
            else:
                current_section.blocks.append(rendered)
                current_section.texts.append(text)
        else:
            if current_section is None:
                continue
            rendered = table_html(block)
            if not rendered:
                continue
            if current_exercise is not None:
                current_exercise["blocks"].append(rendered)
            else:
                current_section.blocks.append(rendered)

    finalize_exercise(current_exercise, current_section)

    return {
        "module_number": module_number,
        "title": title,
        "sections": sections,
        "path": path,
    }


def normalize_list_blocks(blocks: list[str]) -> str:
    html_parts: list[str] = []
    open_list = False
    for block in blocks:
        if block.startswith("<li>"):
            if not open_list:
                html_parts.append("<ul>")
                open_list = True
            html_parts.append(block)
            continue
        if open_list:
            html_parts.append("</ul>")
            open_list = False
        html_parts.append(block)
    if open_list:
        html_parts.append("</ul>")
    return "".join(html_parts).strip()


def derive_description(sections: OrderedDict[str, SectionData]) -> str:
    for key in ("overview", "introduction", "theory"):
        section = sections.get(key)
        if not section:
            continue
        for text in section.texts:
            if len(text) > 40:
                return text[:320].strip()
    return ""


def derive_outcome_title(text: str) -> str:
    clean = normalize_text(text).rstrip(".")
    clean = re.sub(r"^(To|After(?: the)? completion of the course, learners will be able to:)\s+", "", clean, flags=re.IGNORECASE)
    if not clean:
        return "Learning outcome"
    clean = clean[0].upper() + clean[1:]
    words = clean.split()
    if len(words) > 8:
        clean = " ".join(words[:8]).rstrip(",;:") + "..."
    return clean


def build_outcomes(goals_section: SectionData | None) -> list[dict]:
    if not goals_section:
        return []
    outcomes = []
    for text in goals_section.texts:
        lowered = text.lower()
        if lowered.startswith("after the completion") or lowered.startswith("after completion"):
            continue
        if len(text) < 20:
            continue
        outcomes.append(
            {
                "outcome_title": derive_outcome_title(text),
                "outcome_detail": f"<p>{html.escape(text, quote=False)}</p>",
                "outcome_icon": "",
            }
        )
    return outcomes[:10]


def default_hotspots(count: int) -> list[tuple[int, int]]:
    presets = {
        1: [(50, 50)],
        2: [(32, 38), (68, 62)],
        3: [(22, 35), (50, 58), (78, 32)],
        4: [(24, 30), (50, 26), (76, 34), (50, 68)],
        5: [(18, 30), (40, 62), (56, 28), (72, 66), (86, 38)],
    }
    if count in presets:
        return presets[count]
    points = []
    for index in range(count):
        x = 18 + (index % 4) * 22
        y = 28 + ((index // 4) % 3) * 22
        points.append((min(x, 88), min(y, 84)))
    return points


def build_steps(sections: OrderedDict[str, SectionData]) -> list[dict]:
    practice_section = sections.get("practice")
    exercises = practice_section.exercises[:20] if practice_section else []

    if not exercises and practice_section and practice_section.blocks:
        exercises = [
            {
                "title": "Module Activity",
                "content": normalize_list_blocks(practice_section.blocks),
                "texts": practice_section.texts,
            }
        ]

    if not exercises:
        assessment_section = sections.get("assessment")
        if assessment_section and assessment_section.blocks:
            exercises = [
                {
                    "title": "Assessment & Reflection",
                    "content": normalize_list_blocks(assessment_section.blocks),
                    "texts": assessment_section.texts,
                }
            ]

    if not exercises:
        return []

    hotspots = default_hotspots(len(exercises))
    steps = []
    for index, exercise in enumerate(exercises):
        x, y = hotspots[index]
        steps.append(
            {
                "step_title": exercise["title"],
                "step_content": exercise["content"],
                "step_image": "",
                "step_hotspot_x": x,
                "step_hotspot_y": y,
            }
        )
    return steps


def extract_html_tables(content: str) -> list[str]:
    return re.findall(r"<table\b[^>]*>.*?</table>", content, flags=re.IGNORECASE | re.DOTALL)


def extract_table_rows(table_html: str) -> list[list[str]]:
    rows: list[list[str]] = []
    for row_html in re.findall(r"<tr\b[^>]*>(.*?)</tr>", table_html, flags=re.IGNORECASE | re.DOTALL):
        cells = []
        for cell_html in re.findall(r"<(?:td|th)\b[^>]*>(.*?)</(?:td|th)>", row_html, flags=re.IGNORECASE | re.DOTALL):
            text = normalize_text(html.unescape(re.sub(r"<[^>]+>", " ", cell_html)))
            cells.append(text)
        if any(cells):
            rows.append(cells)
    return rows


def parse_assessment_prompt(text: str) -> tuple[int, str] | None:
    match = re.match(r"^Q\s*(\d+)[\.\-:\)]?\s*(.+)$", text, flags=re.IGNORECASE)
    if not match:
        return None
    return int(match.group(1)), normalize_text(match.group(2))


def parse_assessment_option(cells: list[str]) -> tuple[str, str] | None:
    values = [normalize_text(cell) for cell in cells if normalize_text(cell)]
    if not values:
        return None
    if len(values) >= 2 and re.fullmatch(r"[A-Z]", values[0]):
        return values[0], values[1]
    match = re.match(r"^([A-Z])[\)\.\-:]?\s+(.+)$", values[0])
    if match:
        return match.group(1), normalize_text(match.group(2))
    return None


def table_has_answer_prompt(rows: list[list[str]]) -> bool:
    for row in rows:
        for cell in row:
            if "your answer" in normalize_text(cell).lower():
                return True
    return False


def parse_answer_map(answers_section: SectionData | None) -> dict[int, dict[str, str]]:
    answer_map: dict[int, dict[str, str]] = {}
    if not answers_section:
        return answer_map
    for table_html in extract_html_tables(normalize_list_blocks(answers_section.blocks)):
        for row in extract_table_rows(table_html):
            if not row:
                continue
            match = re.match(r"^Q\s*(\d+)(?:\s*[-: ]\s*([A-Z]))?$", row[0], flags=re.IGNORECASE)
            if not match:
                continue
            explanation = normalize_text(row[1]) if len(row) > 1 else ""
            correct_option = (match.group(2) or "").upper()

            if not correct_option and explanation:
                explanation_match = re.match(r"^Correct answer\s*[:\-]?\s*([A-Z])\b\s*(.*)$", explanation, flags=re.IGNORECASE)
                if explanation_match:
                    correct_option = explanation_match.group(1).upper()
                    explanation = normalize_text(explanation_match.group(2))

            if not correct_option:
                continue

            answer_map[int(match.group(1))] = {
                "correct_option": correct_option,
                "explanation": explanation,
            }
    return answer_map


def build_assessment_questions(sections: OrderedDict[str, SectionData]) -> list[dict]:
    assessment_section = sections.get("assessment")
    if not assessment_section or not assessment_section.blocks:
        return []

    answer_map = parse_answer_map(sections.get("answers"))
    questions: list[dict] = []
    current: dict | None = None

    for table_html in extract_html_tables(normalize_list_blocks(assessment_section.blocks)):
        rows = extract_table_rows(table_html)
        if not rows:
            continue

        question_number = None
        prompt = ""
        options: dict[str, str] = {}

        for row in rows:
            if not row:
                continue
            if question_number is None:
                parsed_prompt = parse_assessment_prompt(row[0])
                if parsed_prompt:
                    question_number, prompt = parsed_prompt
                    continue
            parsed_option = parse_assessment_option(row)
            if parsed_option:
                option_key, option_text = parsed_option
                options[option_key] = option_text

        if question_number is not None and prompt:
            if current and current.get("options"):
                answer_entry = answer_map.get(current["number"], {})
                correct_option = answer_entry.get("correct_option", "")
                if correct_option and correct_option in current["options"]:
                    questions.append(
                        {
                            "prompt": current["prompt"],
                            "options": current["options"],
                            "correct_option": correct_option,
                            "explanation": answer_entry.get("explanation", ""),
                        }
                    )
            current = {
                "number": question_number,
                "prompt": prompt,
                "options": options,
            }
            if table_has_answer_prompt(rows):
                answer_entry = answer_map.get(question_number, {})
                correct_option = answer_entry.get("correct_option", "")
                if correct_option and correct_option in options:
                    questions.append(
                        {
                            "prompt": prompt,
                            "options": options,
                            "correct_option": correct_option,
                            "explanation": answer_entry.get("explanation", ""),
                        }
                    )
                current = None
            continue

        if current:
            for option_key, option_text in options.items():
                current["options"][option_key] = option_text
            if table_has_answer_prompt(rows):
                answer_entry = answer_map.get(current["number"], {})
                correct_option = answer_entry.get("correct_option", "")
                if correct_option and correct_option in current["options"]:
                    questions.append(
                        {
                            "prompt": current["prompt"],
                            "options": current["options"],
                            "correct_option": correct_option,
                            "explanation": answer_entry.get("explanation", ""),
                        }
                    )
                current = None

    if current and current.get("options"):
        answer_entry = answer_map.get(current["number"], {})
        correct_option = answer_entry.get("correct_option", "")
        if correct_option and correct_option in current["options"]:
            questions.append(
                {
                    "prompt": current["prompt"],
                    "options": current["options"],
                    "correct_option": correct_option,
                    "explanation": answer_entry.get("explanation", ""),
                }
            )

    return questions[:20]


def build_chapters(sections: OrderedDict[str, SectionData]) -> list[dict]:
    chapters = []
    for key, section in sections.items():
        content = normalize_list_blocks(section.blocks)
        if not content:
            continue
        chapters.append(
            {
                "chapter_title": SECTION_LABELS[key],
                "chapter_content": content,
                "chapter_image": "",
            }
        )
    return chapters[:10]


def select_images(docx_path: Path, module_output_dir: Path, package_root: Path) -> tuple[str, list[str]]:
    cover_path = ""
    gallery_paths: list[str] = []
    candidates = []

    with zipfile.ZipFile(docx_path) as archive:
        media_names = [name for name in archive.namelist() if name.startswith("word/media/")]
        for name in media_names:
            suffix = Path(name).suffix.lower()
            if suffix not in {".png", ".jpg", ".jpeg", ".gif", ".webp"}:
                continue
            data = archive.read(name)
            target_path = module_output_dir / Path(name).name
            target_path.write_bytes(data)

            try:
                with Image.open(target_path) as image:
                    width, height = image.size
            except OSError:
                target_path.unlink(missing_ok=True)
                continue

            if width < 280 or height < 180:
                target_path.unlink(missing_ok=True)
                continue

            score = width * height
            aspect = width / max(height, 1)
            if 1.25 <= aspect <= 2.4 and width >= 900:
                score *= 3
            elif 0.85 <= aspect <= 1.15:
                score *= 0.7

            candidates.append((score, target_path))

    if not candidates:
        return cover_path, gallery_paths

    candidates.sort(key=lambda item: item[0], reverse=True)
    cover_candidate = candidates[0][1]
    cover_path = cover_candidate.relative_to(package_root).as_posix()

    for _, candidate in candidates[1:9]:
        gallery_paths.append(candidate.relative_to(package_root).as_posix())

    return cover_path, gallery_paths


def build_module_payload(parsed_doc: dict, package_root: Path) -> dict:
    module_number = parsed_doc["module_number"]
    sections: OrderedDict[str, SectionData] = parsed_doc["sections"]
    module_dir = package_root / "modules" / f"module-{module_number}"
    media_dir = module_dir / "media"
    media_dir.mkdir(parents=True, exist_ok=True)

    cover_image, gallery_images = select_images(parsed_doc["path"], media_dir, package_root)
    chapters = build_chapters(sections)
    outcomes = build_outcomes(sections.get("goals"))
    steps = build_steps(sections)
    assessment_questions = build_assessment_questions(sections)

    return {
        "module_number": module_number,
        "post_title": parsed_doc["title"],
        "post_name": f"module-{module_number}-{slugify(parsed_doc['title'])}",
        "module_subtitle": MODULE_SUBTITLES[module_number],
        "module_description": derive_description(sections),
        "module_color": MODULE_COLORS[module_number],
        "module_icon": "",
        "module_cover_image": cover_image,
        "module_motto": MODULE_MOTTOS[module_number],
        "module_video_url": "",
        "module_gallery": gallery_images,
        "module_learning_outcomes": outcomes,
        "module_exercise_steps": steps,
        "module_assessment_questions": assessment_questions,
        "module_chapters": chapters,
        "source_document": str(parsed_doc["path"]),
    }


def discover_wp2_docs(source_root: Path) -> list[Path]:
    wp2_dir = source_root / WP2_RELATIVE_DIR
    if not wp2_dir.exists():
        raise FileNotFoundError(f"Could not find WP2 source directory: {wp2_dir}")

    docs = []
    for path in sorted(wp2_dir.glob("*.docx")):
        if path.name.startswith("Module_no._partner"):
            continue
        if path.name.startswith("WP2.4_Design_Learning_Modules_Framework"):
            continue
        docs.append(path)
    if len(docs) != 6:
        raise RuntimeError(f"Expected 6 module documents in {wp2_dir}, found {len(docs)}")
    return docs


def build_manifest(source_root: Path, package_root: Path) -> dict:
    modules = []
    for doc_path in discover_wp2_docs(source_root):
        parsed_doc = parse_wp2_doc(doc_path)
        modules.append(build_module_payload(parsed_doc, package_root))

    modules.sort(key=lambda module: module["module_number"])

    return {
        "schema_version": 2,
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "source": {
            "source_root": str(source_root),
            "primary_source": str(source_root / WP2_RELATIVE_DIR),
        },
        "modules": modules,
    }


def zip_directory(source_dir: Path, output_zip: Path) -> None:
    output_zip.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(output_zip, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path in sorted(source_dir.rglob("*")):
            if path.is_file():
                archive.write(path, path.relative_to(source_dir).as_posix())


def main() -> None:
    args = parse_args()
    source_root = args.source_root.resolve()
    output_zip = args.output.resolve()

    with tempfile.TemporaryDirectory(prefix="wellme-import-package-") as temp_dir_name:
        temp_dir = Path(temp_dir_name)
        package_root = temp_dir / "wellme-import-package"
        package_root.mkdir(parents=True, exist_ok=True)

        manifest = build_manifest(source_root, package_root)
        manifest_path = package_root / "manifest.json"
        manifest_path.write_text(json.dumps(manifest, indent=2, ensure_ascii=False), encoding="utf-8")

        if output_zip.exists():
            output_zip.unlink()

        zip_directory(package_root, output_zip)

    print(f"Created WELLME import package: {output_zip}")
    for module in manifest["modules"]:
        print(f"  Module {module['module_number']}: {module['post_title']}")


if __name__ == "__main__":
    main()
