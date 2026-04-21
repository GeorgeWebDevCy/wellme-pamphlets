from docx import Document
import os, json

base = '/mnt/d/GitHub Projects/wellme-pamphlets/WP 3.1. Creating Hands On Training/'

docs = {
    1: 'WP3.1.Hands on Training Module1_GESEME_WELLME.docx',
    2: 'WP3.1.Hands on Training Module2_EUROPEANPROGRESS_WELLME.docx',
    3: 'WP3_1_Hands_on_Training_Module3_UNIZAR_WELLME_FINAL VERSION.docx',
    4: 'WP3.1.Hands on Training Module4_ETAP_WELLME.docx',
    5: 'WP3.1.Hands on Training Module5_CENTREDOT_WELLME.docx',
    6: 'WP3.1.Hands on Training Module6_Autokreacja.docx',
}

for mod_num, fname in sorted(docs.items()):
    path = os.path.join(base, fname)
    doc = Document(path)
    
    print(f'\n{"="*80}')
    print(f'MODULE {mod_num}: {fname}')
    print(f'Author: {doc.core_properties.author}')
    print(f'{"="*80}')
    
    # Print ALL paragraphs
    print('\n--- PARAGRAPHS ---')
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            print(f'  [{para.style.name}] {text}')
    
    # Print ALL tables with FULL cell content
    for t_idx, table in enumerate(doc.tables):
        print(f'\n--- TABLE {t_idx+1} ({len(table.rows)} rows x {len(table.columns)} cols) ---')
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if cell_text:
                    print(f'\n  [Cell R{r_idx}C{c_idx}]:')
                    # Print line by line
                    for line in cell_text.split('\n'):
                        if line.strip():
                            print(f'    {line}')
