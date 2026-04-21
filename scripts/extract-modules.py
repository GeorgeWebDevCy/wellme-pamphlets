from docx import Document
import json
import os

base = '/mnt/d/GitHub Projects/wellme-pamphlets/WP 3.1. Creating Hands On Training/'

modules = {
    1: 'WP3.1.Hands on Training Module1_GESEME_WELLME.docx',
    2: 'WP3.1.Hands on Training Module2_EUROPEANPROGRESS_WELLME.docx',
    3: 'WP3_1_Hands_on_Training_Module3_UNIZAR_WELLME_FINAL VERSION.docx',
    4: 'WP3.1.Hands on Training Module4_ETAP_WELLME.docx',
    5: 'WP3.1.Hands on Training Module5_CENTREDOT_WELLME.docx',
    6: 'WP3.1.Hands on Training Module6_Autokreacja.docx',
}

module_titles = {
    1: "From Strength To Strength: Positive Psychology For Youth Trainers",
    2: "Bounce Back Stronger",
    3: "Fuel for Flourishing: Healthy Nutrition & Lifestyle in Youth Work",
    4: "Move to Thrive: Integrating Physical Activity into Youth Wellbeing",
    5: "Spaces that Connect: Designing Environments for Youth Belonging",
    6: "Bridges to Adulthood: Guiding Youth Transitions through Community Power",
}

mottos = {
    1: "Savor the Moment — gratitude turns what we have into enough",
    2: "Fall seven times, stand up eight — resilience is a practice",
    3: "Good nutrition fuels not just the body, it fuels confidence, focus, and the courage to dream",
    4: "Walk together, talk together — movement is the language of wellbeing",
    5: "Ο άνθρωπος αγιάζει τον τόπο, κι όχι ο τόπος τον άνθρωπο",
    6: "Growing up is a journey — you don't have to walk it alone",
}

result = {}

for mod_num, fname in modules.items():
    path = os.path.join(base, fname)
    doc = Document(path)

    data = {
        "number": mod_num,
        "title": module_titles[mod_num],
        "motto": mottos[mod_num],
        "author": doc.core_properties.author or "",
        "sections": {},
        "learning_outcomes": "",
        "exercise_title": "",
        "exercise_duration": "",
        "exercise_target": "",
        "exercise_materials": "",
        "exercise_steps": "",
        "tips_for_trainers": "",
        "reflection_questions": "",
        "conclusion_remarks": "",
    }

    # Parse paragraphs for headings
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name == 'Heading 1':
            if 'Introduction' in text:
                data['sections']['introduction_heading'] = text
            elif 'Module Activity' in text:
                data['sections']['activity_heading'] = text
            elif 'Conclusion' in text:
                data['sections']['conclusion_heading'] = text

    # Parse tables
    tables = doc.tables
    for t_idx, table in enumerate(tables):
        cell_texts = []
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    t = para.text.strip()
                    if t:
                        cell_texts.append(t)

        full_text = "\n".join(cell_texts)

        if t_idx == 0:
            data['sections']['theoretical_background'] = full_text
        elif t_idx == 1:
            data['learning_outcomes'] = full_text
        elif t_idx == 2:
            data['sections']['exercise'] = full_text
            for i, line in enumerate(cell_texts):
                if 'Title of the Exercise' in line:
                    if i + 1 < len(cell_texts):
                        data['exercise_title'] = cell_texts[i + 1]
                if 'Duration:' in line:
                    duration_parts = []
                    for j in range(i+1, min(i+5, len(cell_texts))):
                        if cell_texts[j].startswith('Target') or cell_texts[j].startswith('Materials'):
                            break
                        duration_parts.append(cell_texts[j])
                    if duration_parts:
                        data['exercise_duration'] = " ".join(duration_parts)
                if 'Materials Needed:' in line:
                    mat_parts = []
                    for j in range(i+1, min(i+10, len(cell_texts))):
                        if cell_texts[j].startswith('Step-by-Step') or cell_texts[j].startswith('Tips'):
                            break
                        mat_parts.append(cell_texts[j])
                    if mat_parts:
                        data['exercise_materials'] = "\n".join(mat_parts)
                if 'Step-by-Step Implementation' in line:
                    step_parts = []
                    for j in range(i+1, len(cell_texts)):
                        if cell_texts[j].startswith('Tips for Trainers'):
                            break
                        step_parts.append(cell_texts[j])
                    if step_parts:
                        data['exercise_steps'] = "\n".join(step_parts)
                if 'Tips for Trainers' in line:
                    tip_parts = []
                    for j in range(i+1, len(cell_texts)):
                        if cell_texts[j].startswith('---'):
                            break
                        tip_parts.append(cell_texts[j])
                    if tip_parts:
                        data['tips_for_trainers'] = "\n".join(tip_parts)
        elif t_idx == 3:
            data['sections']['conclusion'] = full_text
            questions = [l for l in cell_texts if l.endswith('?')]
            data['reflection_questions'] = "\n".join(questions)

    result[mod_num] = data
    print(f"Module {mod_num} ({module_titles[mod_num]}): OK")

out_path = '/mnt/d/GitHub Projects/wellme-pamphlets/scripts/module-data.json'
os.makedirs(os.path.dirname(out_path), exist_ok=True)
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump(result, f, ensure_ascii=False, indent=2)

print(f"\nExtracted to {out_path}")
print(f"Modules: {len(result)}")
